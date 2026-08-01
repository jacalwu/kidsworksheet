"""
檔案解析模組：解析 PDF 與 Word (.docx) 檔案，擷取文字內容並結構化為 JSON。

支援功能：
- PDF 文字擷取（使用 PyPDF2 與 pdfplumber）
- PDF 圖片 OCR（使用 pytesseract + pdf2image，需安裝 Tesseract 與 Poppler）
- Word (.docx) 文字擷取（使用 python-docx）
- 將擷取內容結構化為標準 JSON 格式
"""

import os
import json
import hashlib
from datetime import datetime, timezone
from typing import Optional
from io import BytesIO


# ── PDF 解析 ──────────────────────────────────────────────

def extract_pdf_text(file_bytes: bytes, use_ocr: bool = False) -> dict:
    """從 PDF 位元組中擷取文字內容

    Args:
        file_bytes: PDF 檔案的位元組內容
        use_ocr: 是否啟用 OCR 掃描圖片中的文字

    Returns:
        dict: 結構化 JSON，包含 metadata 與內容
    """
    result = {
        "metadata": {
            "file_type": "pdf",
            "extraction_method": "text",
            "ocr_enabled": use_ocr,
            "pages": 0,
            "processed_at": datetime.now(timezone.utc).isoformat(),
        },
        "content": [],
    }

    # 方法 1：使用 pdfplumber（排版保留較佳）
    try:
        import pdfplumber

        with pdfplumber.open(BytesIO(file_bytes)) as pdf:
            result["metadata"]["pages"] = len(pdf.pages)
            for i, page in enumerate(pdf.pages):
                page_data = {
                    "page_number": i + 1,
                    "text": "",
                    "tables": [],
                    "ocr_text": "",
                }

                # 擷取文字
                text = page.extract_text()
                if text:
                    page_data["text"] = text.strip()

                # 擷取表格
                tables = page.extract_tables()
                if tables:
                    page_data["tables"] = [
                        [[cell or "" for cell in row] for row in table]
                        for table in tables
                    ]

                # OCR（若啟用，且頁面文字很少）
                if use_ocr and (not text or len(text.strip()) < 50):
                    ocr_text = _ocr_pdf_page(file_bytes, i)
                    if ocr_text:
                        page_data["ocr_text"] = ocr_text
                        result["metadata"]["extraction_method"] = "text+ocr"

                result["content"].append(page_data)

        if result["content"]:
            return result
    except ImportError:
        pass  # 繼續嘗試下一個方法
    except Exception as e:
        # pdfplumber 失敗，嘗試 PyPDF2
        pass

    # 方法 2：使用 PyPDF2（備援方案）
    try:
        from PyPDF2 import PdfReader

        reader = PdfReader(BytesIO(file_bytes))
        result["metadata"]["pages"] = len(reader.pages)

        for i, page in enumerate(reader.pages):
            page_data = {
                "page_number": i + 1,
                "text": "",
                "tables": [],
                "ocr_text": "",
            }

            text = page.extract_text()
            if text:
                page_data["text"] = text.strip()

            # OCR（若啟用）
            if use_ocr and (not text or len(text.strip()) < 50):
                ocr_text = _ocr_pdf_page(file_bytes, i)
                if ocr_text:
                    page_data["ocr_text"] = ocr_text
                    result["metadata"]["extraction_method"] = "text+ocr"

            result["content"].append(page_data)

        return result
    except Exception as e:
        # 兩種方法都失敗
        result["metadata"]["error"] = str(e)
        result["metadata"]["extraction_method"] = "failed"
        return result


def _ocr_pdf_page(file_bytes: bytes, page_index: int) -> str:
    """對 PDF 的指定頁面執行 OCR

    Args:
        file_bytes: PDF 檔案的位元組內容
        page_index: 頁面索引（0-based）

    Returns:
        str: OCR 辨識的文字，若失敗則回傳空字串
    """
    try:
        from pdf2image import convert_from_bytes
        import pytesseract

        images = convert_from_bytes(
            file_bytes, first_page=page_index + 1, last_page=page_index + 1, dpi=200
        )
        if images:
            text = pytesseract.image_to_string(images[0], lang="chi_tra+eng")
            return text.strip()
    except ImportError:
        pass  # 缺少依賴
    except Exception:
        pass  # OCR 失敗

    return ""


def get_pdf_page_count(file_bytes: bytes) -> int:
    """取得 PDF 頁數"""
    try:
        import pdfplumber

        with pdfplumber.open(BytesIO(file_bytes)) as pdf:
            return len(pdf.pages)
    except Exception:
        try:
            from PyPDF2 import PdfReader

            reader = PdfReader(BytesIO(file_bytes))
            return len(reader.pages)
        except Exception:
            return 0


# ── Word (.docx) 解析 ─────────────────────────────────────

def extract_docx_text(file_bytes: bytes) -> dict:
    """從 Word (.docx) 位元組中擷取文字內容

    Args:
        file_bytes: .docx 檔案的位元組內容

    Returns:
        dict: 結構化 JSON，包含 metadata 與內容
    """
    result = {
        "metadata": {
            "file_type": "docx",
            "extraction_method": "python-docx",
            "paragraphs": 0,
            "processed_at": datetime.now(timezone.utc).isoformat(),
        },
        "content": [],
    }

    try:
        from docx import Document

        doc = Document(BytesIO(file_bytes))
        paragraphs = []

        current_section = {"heading": "", "heading_level": 0, "paragraphs": []}

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                # 空行表示章節邊界
                if current_section["paragraphs"]:
                    paragraphs.append(dict(current_section))
                    current_section = {"heading": "", "heading_level": 0, "paragraphs": []}
                continue

            # 判斷是否為標題
            if para.style.name.startswith("Heading") or para.style.name.startswith("heading"):
                # 儲存上一個章節
                if current_section["paragraphs"] or current_section["heading"]:
                    paragraphs.append(dict(current_section))

                level = 1
                try:
                    level_str = para.style.name.replace("Heading", "").replace("heading", "").strip()
                    level = int(level_str) if level_str else 1
                except ValueError:
                    level = 1

                current_section = {
                    "heading": text,
                    "heading_level": level,
                    "paragraphs": [],
                }
            else:
                current_section["paragraphs"].append(text)

        # 儲存最後一個章節
        if current_section["paragraphs"] or current_section["heading"]:
            paragraphs.append(dict(current_section))

        # 若無標題結構，將所有段落放在一個區塊中
        if not paragraphs:
            all_text = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            if all_text:
                paragraphs = [{"heading": "", "heading_level": 0, "paragraphs": all_text}]

        result["content"] = paragraphs
        result["metadata"]["paragraphs"] = sum(
            len(s["paragraphs"]) for s in paragraphs
        )

        return result
    except Exception as e:
        result["metadata"]["error"] = str(e)
        result["metadata"]["extraction_method"] = "failed"
        return result


# ── 整合解析介面 ──────────────────────────────────────────

def parse_file(file_bytes: bytes, file_name: str, use_ocr: bool = False) -> dict:
    """根據副檔名自動選擇解析器

    Args:
        file_bytes: 檔案位元組內容
        file_name: 原始檔名（用於判斷副檔名）
        use_ocr: 是否對 PDF 啟用 OCR

    Returns:
        dict: 結構化 JSON 結果
    """
    ext = os.path.splitext(file_name)[1].lower()

    if ext == ".pdf":
        result = extract_pdf_text(file_bytes, use_ocr=use_ocr)
    elif ext in (".docx", ".doc"):
        result = extract_docx_text(file_bytes)
    else:
        return {
            "metadata": {
                "file_type": ext,
                "error": f"不支援的檔案類型：{ext}",
                "processed_at": datetime.now(timezone.utc).isoformat(),
            },
            "content": [],
        }

    # 加入檔案層級 metadata
    result["metadata"]["file_name"] = file_name
    result["metadata"]["file_size_bytes"] = len(file_bytes)
    result["metadata"]["file_hash"] = hashlib.sha256(file_bytes).hexdigest()[:16]

    return result


def format_parsed_content(parsed: dict) -> str:
    """將解析結果格式化為純文字，便於傳遞給 LLM

    Args:
        parsed: parse_file 回傳的結構化 JSON

    Returns:
        str: 格式化的純文字內容
    """
    lines = []
    file_name = parsed.get("metadata", {}).get("file_name", "未知檔案")
    lines.append(f"=== 檔案：{file_name} ===")

    for section in parsed.get("content", []):
        if isinstance(section, dict):
            # PDF 頁面格式
            if "page_number" in section:
                lines.append(f"\n--- 第 {section['page_number']} 頁 ---")
                if section.get("text"):
                    lines.append(section["text"])
                if section.get("ocr_text"):
                    lines.append(f"[OCR 辨識]: {section['ocr_text']}")
                if section.get("tables"):
                    for ti, table in enumerate(section["tables"]):
                        lines.append(f"\n[表格 {ti + 1}]")
                        for row in table:
                            lines.append(" | ".join(str(c) for c in row))
            # Word 章節格式
            else:
                heading = section.get("heading", "")
                if heading:
                    prefix = "#" * min(section.get("heading_level", 1), 4)
                    lines.append(f"\n{prefix} {heading}")
                for p in section.get("paragraphs", []):
                    lines.append(p)

    return "\n".join(lines)


def get_parsed_json_for_display(parsed: dict) -> str:
    """將解析結果轉為可顯示的 JSON 字串（用於 UI 展示）"""
    # 簡化 content 以利顯示
    display = {
        "metadata": parsed.get("metadata", {}),
        "content_summary": [],
    }
    for section in parsed.get("content", []):
        if isinstance(section, dict):
            summary = {}
            if "page_number" in section:
                summary["page"] = section["page_number"]
                summary["text_length"] = len(section.get("text", ""))
                summary["has_ocr"] = bool(section.get("ocr_text"))
                summary["table_count"] = len(section.get("tables", []))
            else:
                summary["heading"] = section.get("heading", "")
                summary["paragraph_count"] = len(section.get("paragraphs", []))
            display["content_summary"].append(summary)

    return json.dumps(display, ensure_ascii=False, indent=2)
