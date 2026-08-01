"""
生成模組：使用 DeepSeek LLM 與 Web Search 產生 worksheet 或考試題目，
並輸出為 Word (.docx) 格式供下載。

支援兩種模式：
- cloud：使用真實 DeepSeek API + Web Search
- local（mock）：回傳模擬資料，不需 API key
"""

import os
import json
from datetime import datetime, timezone
from typing import Optional
from io import BytesIO


# ── 設定載入 ──────────────────────────────────────────────

def load_config() -> dict:
    """從資料庫、TOML 設定檔或 Streamlit secrets 載入設定

    優先順序：DB 設定 > Streamlit secrets > config.toml > 環境變數 > 預設值
    """
    config = {
        "LEARNING_DEPLOYMENT": "local",
        "LEARNING_API_KEY": "",
        "LEARNING_LLM_MODEL": "deepseek-chat",
        "LEARNING_LLM_URL": "https://api.deepseek.com/v1",
        "LEARNING_WEB_SEARCH_API_KEY": "",
        "LEARNING_WEB_SEARCH_ENGINE": "serpapi",
    }

    # 0. 嘗試從資料庫讀取（最高優先，管理員可透過 UI 修改）
    try:
        from utils.database import get_all_settings
        db_settings = get_all_settings()
        for key in config:
            if key in db_settings and db_settings[key]:
                config[key] = db_settings[key]
    except Exception:
        pass  # DB 尚未初始化或無法連線，使用其他來源

    # 1. 嘗試從環境變數讀取
    for key in config:
        env_val = os.environ.get(key)
        if env_val:
            config[key] = env_val

    # 2. 嘗試從 config.toml 讀取
    try:
        import tomllib
    except ImportError:
        import tomli as tomllib

    try:
        with open("config.toml", "rb") as f:
            toml_cfg = tomllib.load(f)
        for key in config:
            if key in toml_cfg:
                config[key] = toml_cfg[key]
    except FileNotFoundError:
        pass
    except Exception:
        pass

    # 3. 嘗試從 Streamlit secrets 讀取（優先權最高）
    try:
        import streamlit as st
        for key in config:
            try:
                val = st.secrets.get(key)
                if val:
                    config[key] = val
            except Exception:
                pass
    except Exception:
        pass

    return config


# ── LLM 呼叫 ──────────────────────────────────────────────

def call_llm(
    messages: list[dict],
    config: Optional[dict] = None,
) -> str:
    """呼叫 DeepSeek LLM API（OpenAI 相容介面）

    Args:
        messages: 聊天訊息列表
        config: 設定字典（若為 None 則自動載入）

    Returns:
        str: LLM 回覆的文字內容
    """
    if config is None:
        config = load_config()

    # Mock 模式：回傳模擬資料
    if config.get("LEARNING_DEPLOYMENT", "local") == "local":
        return _mock_llm_response(messages)

    # Cloud 模式：呼叫真實 API
    try:
        from openai import OpenAI

        client = OpenAI(
            api_key=config["LEARNING_API_KEY"],
            base_url=config.get("LEARNING_LLM_URL", "https://api.deepseek.com/v1"),
        )

        response = client.chat.completions.create(
            model=config.get("LEARNING_LLM_MODEL", "deepseek-chat"),
            messages=messages,
            temperature=0.7,
            max_tokens=4096,
        )

        return response.choices[0].message.content
    except ImportError:
        return "⚠️ 請安裝 openai 套件：pip install openai"
    except Exception as e:
        return f"⚠️ LLM API 呼叫失敗：{str(e)}"


def _mock_llm_response(messages: list[dict]) -> str:
    """Mock 模式：根據使用者 prompt 回傳模擬的題目內容"""
    last_message = messages[-1]["content"] if messages else ""

    # 判斷是 worksheet 還是 exam
    if "worksheet" in last_message.lower() or "練習" in last_message:
        return _mock_worksheet(last_message)
    else:
        return _mock_exam(last_message)


def _mock_worksheet(prompt: str) -> str:
    """產生模擬的 worksheet 內容"""
    return """# 復習練習題（Worksheet）

## 第一部分：選擇題

1. 下列哪一項是正確的？
   A) 選項一
   B) 選項二
   C) 選項三
   D) 選項四

2. 根據文章內容，主要討論的主題是什麼？
   A) 科技發展
   B) 環境保護
   C) 教育改革
   D) 文化交流

## 第二部分：填空題

1. 文章中提到的核心理念是 __________。
2. 作者認為最重要的因素是 __________。

## 第三部分：簡答題

1. 請簡述文章的主要論點。（50-100 字）

---

> 📝 此為 Mock 模式產生的範例內容。
> 將 LEARNING_DEPLOYMENT 設為 "cloud" 並設定 API Key 以使用真實 LLM 生成。
"""


def _mock_exam(prompt: str) -> str:
    """產生模擬的考試題目內容"""
    return """# 模擬考試卷

**考試時間：60 分鐘** | **總分：100 分**

---

## 一、選擇題（每題 5 分，共 25 分）

1. 下列敘述何者正確？
   A) 選項一
   B) 選項二
   C) 選項三
   D) 選項四

2. 文章中作者的主要立場是？
   A) 支持
   B) 反對
   C) 中立
   D) 未表明

3-5. （略）

---

## 二、填空題（每題 5 分，共 25 分）

1. __________ 是本文的核心概念。
2. 作者在第三段提到 __________ 的重要性。

---

## 三、申論題（每題 25 分，共 50 分）

1. 請分析文章的主要觀點，並提出你的看法。（200-300 字）

2. 請比較文章中提到的兩種不同觀點，並說明你支持哪一方及其理由。

---

> 📝 此為 Mock 模式產生的範例內容。
> 將 LEARNING_DEPLOYMENT 設為 "cloud" 並設定 API Key 以使用真實 LLM 生成。
"""


# ── Web Search ─────────────────────────────────────────────

def web_search(query: str, config: Optional[dict] = None) -> list[dict]:
    """執行網路搜尋

    Args:
        query: 搜尋查詢字串
        config: 設定字典

    Returns:
        list[dict]: 搜尋結果列表，每項包含 title, link, snippet
    """
    if config is None:
        config = load_config()

    # Mock 模式
    if config.get("LEARNING_DEPLOYMENT", "local") == "local":
        return _mock_web_search(query)

    engine = config.get("LEARNING_WEB_SEARCH_ENGINE", "serpapi")
    api_key = config.get("LEARNING_WEB_SEARCH_API_KEY", "")

    if not api_key:
        return [{"title": "未設定搜尋 API Key", "link": "", "snippet": "請在 config.toml 中設定 LEARNING_WEB_SEARCH_API_KEY"}]

    if engine == "serpapi":
        return _serpapi_search(query, api_key)
    elif engine == "bing":
        return _bing_search(query, api_key)
    else:
        return [{"title": "不支援的搜尋引擎", "link": "", "snippet": f"引擎類型：{engine}"}]


def _mock_web_search(query: str) -> list[dict]:
    """Mock 搜尋結果"""
    return [
        {
            "title": f"[Mock] 搜尋結果：{query[:50]}",
            "link": "https://example.com/result1",
            "snippet": "這是一個模擬的搜尋結果。在 Mock 模式下不會執行真實搜尋。",
        },
        {
            "title": "[Mock] 相關資料",
            "link": "https://example.com/result2",
            "snippet": "設定 LEARNING_DEPLOYMENT=\"cloud\" 並提供 Web Search API Key 以獲得真實搜尋結果。",
        },
    ]


def _serpapi_search(query: str, api_key: str) -> list[dict]:
    """使用 SerpAPI 執行搜尋"""
    try:
        import requests
        params = {
            "q": query,
            "api_key": api_key,
            "engine": "google",
            "num": 5,
        }
        resp = requests.get("https://serpapi.com/search", params=params, timeout=10)
        data = resp.json()

        results = []
        for item in data.get("organic_results", [])[:5]:
            results.append({
                "title": item.get("title", ""),
                "link": item.get("link", ""),
                "snippet": item.get("snippet", ""),
            })
        return results or _mock_web_search(query)
    except Exception as e:
        return [{"title": "搜尋失敗", "link": "", "snippet": str(e)}]


def _bing_search(query: str, api_key: str) -> list[dict]:
    """使用 Bing Search API 執行搜尋"""
    try:
        import requests
        headers = {"Ocp-Apim-Subscription-Key": api_key}
        params = {"q": query, "count": 5, "mkt": "zh-TW"}
        resp = requests.get(
            "https://api.bing.microsoft.com/v7.0/search",
            headers=headers,
            params=params,
            timeout=10,
        )
        data = resp.json()

        results = []
        for item in data.get("webPages", {}).get("value", [])[:5]:
            results.append({
                "title": item.get("name", ""),
                "link": item.get("url", ""),
                "snippet": item.get("snippet", ""),
            })
        return results or _mock_web_search(query)
    except Exception as e:
        return [{"title": "搜尋失敗", "link": "", "snippet": str(e)}]


# ── 生成 Worksheet / Exam ──────────────────────────────────

def build_generation_prompt(
    content_text: str,
    kid_name: str,
    grade: str,
    generate_type: str,
    search_results: Optional[list[dict]] = None,
) -> list[dict]:
    """建立傳送給 LLM 的 prompt

    Args:
        content_text: 擷取後的檔案文字內容
        kid_name: Kid 名稱
        grade: 年級
        generate_type: "worksheet" 或 "exam"
        search_results: 網路搜尋結果（可選）

    Returns:
        list[dict]: 符合 OpenAI Chat API 格式的 messages
    """
    type_label = "練習題（Worksheet）" if generate_type == "worksheet" else "考試題目（Exam）"
    type_instruction = (
        "設計一份復習練習題（worksheet），包含選擇題、填充題、簡答題，並提供参考答案。"
        if generate_type == "worksheet"
        else "設計一份正式考試卷（exam），包含選擇題、填充題、申論題，標註配分與考試時間，並提供参考答案與評分標準。"
    )

    search_context = ""
    if search_results:
        search_context = "\n\n【補充搜尋資料】\n"
        for i, r in enumerate(search_results[:5]):
            search_context += f"\n{i + 1}. {r['title']}\n   {r['snippet']}\n"

    system_prompt = f"""你是一位專業的教育內容設計師，擅長為不同年級的學生設計復習資料。

你的任務是：{type_instruction}

設計原則：
- 題目難度需適合 {grade} 年級學生的程度
- 題目應涵蓋復習資料中的重要概念
- 題型多樣化，包含不同難度層次
- 使用繁體中文出題
- 格式清晰，方便列印使用
- 最後附上完整的参考答案

輸出格式為 Markdown，包含清楚的標題與分節。"""

    user_prompt = f"""請根據以下復習資料，為 {kid_name}（{grade} 年級）{type_label}。

【復習資料內容】
{content_text[:8000]}
{search_context}

請生成完整的{type_label}。"""

    return [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt},
    ]


def generate_document(
    content_text: str,
    kid_name: str,
    grade: str,
    generate_type: str,
    use_search: bool = False,
    config: Optional[dict] = None,
) -> tuple[str, str]:
    """生成 worksheet 或 exam 文件

    Args:
        content_text: 擷取後的檔案文字內容
        kid_name: Kid 名稱
        grade: 年級
        generate_type: "worksheet" 或 "exam"
        use_search: 是否使用網路搜尋補充資料
        config: 設定字典

    Returns:
        tuple[str, str]: (markdown 內容, 使用的 LLM 模型名稱)
    """
    if config is None:
        config = load_config()

    # 執行網路搜尋（若啟用）
    search_results = None
    if use_search:
        search_query = f"{grade}年級 復習資料 題目"
        search_results = web_search(search_query, config)

    # 建立 prompt
    messages = build_generation_prompt(
        content_text, kid_name, grade, generate_type, search_results
    )

    # 呼叫 LLM
    llm_response = call_llm(messages, config)

    model = config.get("LEARNING_LLM_MODEL", "mock-mode")
    return llm_response, model


# ── DOCX 輸出 ─────────────────────────────────────────────

def markdown_to_docx_bytes(markdown_content: str, title: str = "復習資料") -> bytes:
    """將 Markdown 內容轉換為 .docx 檔案位元組

    Args:
        markdown_content: Markdown 格式的內容
        title: 文件標題

    Returns:
        bytes: .docx 檔案的位元組內容
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # 設定預設字型
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Microsoft JhengHei"
        font.size = Pt(12)

        # 解析簡易 Markdown 並加入文件
        lines = markdown_content.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i].strip()

            if not line:
                i += 1
                continue

            # 標題 H1
            if line.startswith("# ") and not line.startswith("## "):
                p = doc.add_heading(line[2:], level=1)
                i += 1
                continue

            # 標題 H2
            if line.startswith("## "):
                p = doc.add_heading(line[3:], level=2)
                i += 1
                continue

            # 標題 H3
            if line.startswith("### "):
                p = doc.add_heading(line[4:], level=3)
                i += 1
                continue

            # 分隔線
            if line.startswith("---"):
                doc.add_paragraph("─" * 40)
                i += 1
                continue

            # 引用
            if line.startswith("> "):
                p = doc.add_paragraph()
                run = p.add_run(line[2:])
                run.italic = True
                run.font.color.rgb = RGBColor(100, 100, 100)
                i += 1
                continue

            # 列表
            if line.startswith("- ") or line.startswith("* "):
                p = doc.add_paragraph(line[2:], style="List Bullet")
                i += 1
                continue

            # 編號列表（簡易判斷：數字開頭 + 點）
            if line and line[0].isdigit() and ". " in line[:5]:
                text = line.split(". ", 1)[1] if ". " in line else line
                p = doc.add_paragraph(text, style="List Number")
                i += 1
                continue

            # 粗體標題行（以 ** 開頭）
            if line.startswith("**") and line.endswith("**"):
                p = doc.add_paragraph()
                run = p.add_run(line[2:-2])
                run.bold = True
                i += 1
                continue

            # 一般段落
            p = doc.add_paragraph(line)
            i += 1

        # 加入頁尾
        doc.add_paragraph("\n─" * 30)
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run(
            f"生成時間：{datetime.now().strftime('%Y-%m-%d %H:%M')}  |  AI 學習助手"
        )
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(150, 150, 150)

        # 儲存到 BytesIO
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output.getvalue()
    except ImportError:
        # 若 python-docx 未安裝，回退為純文字
        fallback = f"{title}\n\n{markdown_content}"
        return fallback.encode("utf-8")
    except Exception as e:
        error_msg = f"生成 DOCX 時發生錯誤：{str(e)}\n\n{markdown_content}"
        return error_msg.encode("utf-8")


def create_docx_download(
    markdown_content: str,
    generate_type: str,
    kid_name: str,
) -> tuple[bytes, str]:
    """建立可下載的 .docx 檔案

    Args:
        markdown_content: Markdown 格式的內容
        generate_type: "worksheet" 或 "exam"
        kid_name: Kid 名稱

    Returns:
        tuple[bytes, str]: (檔案位元組, 建議檔名)
    """
    type_label = "練習題" if generate_type == "worksheet" else "考試卷"
    title = f"{kid_name}_{type_label}"
    filename = f"{kid_name}_{type_label}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"

    docx_bytes = markdown_to_docx_bytes(markdown_content, title)
    return docx_bytes, filename


def is_real_generation(config: Optional[dict] = None) -> bool:
    """判斷目前是否為真實生成模式（非 mock）"""
    if config is None:
        config = load_config()
    return config.get("LEARNING_DEPLOYMENT", "local") == "cloud"
